"""
模块描述：Word 文档读取与批注工具，负责段落定位并生成带批注的 docx 副本。
"""

import os
from os.path import exists

from docx import Document
from docx.shared import Inches
import docx


class WordAnnotator:
    def __init__(self, file_path):
        if not exists(file_path):
            print(f"未找到文件{file_path}")
            return
        self.file_path = file_path
        self.doc = Document(file_path)
    def reader(self):
        """读取文本，返回json格式输出"""
        word_list = []
        for i, paragraph in enumerate(self.doc.paragraphs):
            if not paragraph.text:
                continue

            para = {
                "index": i,
                "text": paragraph.text,
            }
            word_list.append(para)
        return word_list
    def writer(self, index, text, output_path=None):
        if output_path is None:
            base_name, ext = os.path.splitext(self.file_path)
            if not base_name.endswith('_lawyance'):
                base_name = f"{base_name}_lawyance"
            output_path = f"{base_name}{ext}"

        if not exists(output_path):
            doc = Document(self.file_path)
        else:
            doc = Document(output_path)

        print(f"正在批注段落 {index}……")

        # 范围检查
        if index < 0 or index >= len(doc.paragraphs):
            print(f"错误：索引 {index} 超出范围 (总段落数: {len(doc.paragraphs)})")
            return False

        target_para = doc.paragraphs[index]

        # 确保段落有 runs，否则 add_comment 可能会失败
        if not target_para.runs:
            target_para.add_run("")

        try:
            # 注意：标准 python-docx 可能不支持 add_comment，这里假设环境已提供此功能
            comment = doc.add_comment(
                runs=target_para.runs,
                text=text,
                author="工大法智",
                initials="lawyance",
            )
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"批注写入失败: {e}")
            return False

def word_reader(file_path):
    try:
        word = WordAnnotator(file_path)
        if not hasattr(word, 'doc'):
            return []
        words = word.reader()
        return words
    except Exception as e:
        print(f"读取 Word 失败: {e}")
        return []

def word_writer(file_path, index, text, output_path=None):
    try:
        word = WordAnnotator(file_path)
        if not hasattr(word, 'doc'):
            return False
        return word.writer(index, text, output_path)
    except Exception as e:
        print(f"写入 Word 失败: {e}")
        return False

# 最简测试用例
if __name__ == "__main__":
    words = word_reader("sample.docx")
    print(words)
    # words = word.writer(10, "这是一只inu")
    word_writer("sample.docx",10, "测试你是不是吗喽")
